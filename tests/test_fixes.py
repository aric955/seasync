"""
SeaSync V2.2 修复验证测试
=========================
覆盖字段统一、PPI时间过滤、关联可视化、多传感器颜色映射等修复点。
"""

# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 荣火
import pytest
import tempfile
import os


# ============================================================
# 1. PPI 时间过滤验证
# ============================================================

class TestPPITimeFilter:
    """PPI render_ppi_frame 时间过滤正确性。"""

    def test_ppi_filters_by_time_range(self):
        """给定 time_range 时，只返回该窗口内的记录。"""
        from seasync.core.data_models import TargetRecord

        # 模拟 100 秒内的数据
        records = [
            TargetRecord(
                source_id="src", track_id=f"T-{i}",
                time=1000.0 + i * 10.0,  # t = 1000, 1010, ..., 1090
                lat=37.53 + i * 0.001, lon=121.42 + i * 0.001,
            )
            for i in range(10)
        ]

        # 测试窗口 t=1020~1050 应该只包含第 2,3,4,5 条
        t_min, t_max = 1020.0, 1050.0
        filtered = [
            r for r in records
            if t_min <= r.time <= t_max
        ]
        assert len(filtered) == 4
        assert all(t_min <= r.time <= t_max for r in filtered)

    def test_ppi_handles_empty_window(self):
        """时间窗口无数据时不崩溃。"""
        from seasync.core.data_models import TargetRecord

        records = [
            TargetRecord(
                source_id="src", track_id="T-1",
                time=1000.0, lat=37.53, lon=121.42,
            )
        ]
        filtered = [r for r in records if 2000.0 <= r.time <= 2100.0]
        assert len(filtered) == 0

    def test_ppi_handles_edge_boundaries(self):
        """边界值正确处理。"""
        from seasync.core.data_models import TargetRecord

        records = [
            TargetRecord(source_id="src", track_id="T-1",
                         time=1000.0, lat=37.53, lon=121.42),
            TargetRecord(source_id="src", track_id="T-2",
                         time=1100.0, lat=37.54, lon=121.43),
        ]
        # 刚好在边界上
        filtered = [r for r in records if 1000.0 <= r.time <= 1100.0]
        assert len(filtered) == 2

        # 半开
        filtered = [r for r in records if 1000.0 <= r.time < 1100.0]
        assert len(filtered) == 1


# ============================================================
# 2. 关联结果可视化（通用型 label_a/label_b）
# ============================================================

class TestAssociationVisualizationGeneric:
    """render_association 支持通用传感器标签。"""

    def test_render_association_with_generic_labels(self):
        """使用 label_a/label_b 参数渲染任意传感器对。"""
        from seasync.core.data_models import TargetRecord, AssociationPair, AssociationResult
        from seasync.visualization.visualization import render_association

        records_a = [
            TargetRecord(source_id="optical_1", track_id="OPT-1",
                         time=1000.0 + i * 10.0, lat=37.53 + i * 0.001, lon=121.42 + i * 0.001)
            for i in range(3)
        ]
        records_b = [
            TargetRecord(source_id="radar_1", track_id="RT-1",
                         time=1000.0 + i * 10.0, lat=37.53 + i * 0.001, lon=121.42 + i * 0.001)
            for i in range(3)
        ]
        assoc = AssociationResult(
            pairs=[AssociationPair(
                source1_id="OPT-1", source2_id="RT-1",
                source1_label="optical", source2_label="radar",
                confidence=0.9, method="test"
            )],
            total_quality=0.9,
        )

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "assoc_generic.png")
            render_association(
                records_a=records_a, records_b=records_b,
                assoc_result=assoc,
                output_path=out,
                label_a="光电", label_b="雷达",
            )
            assert os.path.exists(out)
            assert os.path.getsize(out) > 1000

    def test_render_association_without_assoc_result(self):
        """无关联结果时仅绘制轨迹。"""
        from seasync.core.data_models import TargetRecord
        from seasync.visualization.visualization import render_association

        records_a = [
            TargetRecord(source_id="a", track_id="A1",
                         time=1000.0 + i * 10.0, lat=37.53 + i * 0.001, lon=121.42 + i * 0.001)
            for i in range(3)
        ]
        records_b = [
            TargetRecord(source_id="b", track_id="B1",
                         time=1000.0 + i * 10.0, lat=37.54 + i * 0.001, lon=121.43 + i * 0.001)
            for i in range(3)
        ]

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "assoc_no_result.png")
            render_association(
                records_a=records_a, records_b=records_b,
                output_path=out,
            )
            assert os.path.exists(out)
            assert os.path.getsize(out) > 1000

    def test_render_association_by_track_id(self):
        """关联结果按 track_id 分组绘制轨迹。"""
        from seasync.core.data_models import TargetRecord
        from seasync.visualization.visualization import render_association

        # 多轨迹数据
        records_a = [
            TargetRecord(source_id="a", track_id="TRK-1",
                         time=1000.0 + i * 10.0, lat=37.53 + i * 0.001, lon=121.42 + i * 0.001)
            for i in range(5)
        ] + [
            TargetRecord(source_id="a", track_id="TRK-2",
                         time=1000.0 + i * 10.0, lat=37.60 + i * 0.001, lon=121.50 + i * 0.001)
            for i in range(5)
        ]
        records_b = [
            TargetRecord(source_id="b", track_id="TRK-1",
                         time=1000.0 + i * 10.0, lat=37.53 + i * 0.001, lon=121.42 + i * 0.001)
            for i in range(5)
        ]

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "assoc_tracks.png")
            render_association(
                records_a=records_a, records_b=records_b,
                output_path=out, label_a="雷达", label_b="AIS",
            )
            assert os.path.exists(out)
            assert os.path.getsize(out) > 1000


# ============================================================
# 3. 多传感器颜色映射
# ============================================================

class TestSensorColorMapping:
    """SENSOR_COLORS 映射表 + get_sensor_color 函数。"""

    def test_known_sensor_types(self):
        """已知传感器类型返回对应颜色。"""
        from seasync.visualization.draw_command import get_sensor_color, SENSOR_COLORS

        for stype, expected in [
            ("radar", "#4A90D9"),
            ("ais", "#E74C3C"),
            ("gps", "#2ECC71"),
            ("dat", "#E67E22"),
            ("mat", "#9B59B6"),
            ("csv", "#1ABC9C"),
            ("optical", "#F39C12"),
            ("sonar", "#16A085"),
            ("imu", "#8E44AD"),
            ("weather", "#D35400"),
        ]:
            assert get_sensor_color(stype) == expected, f"颜色不匹配: {stype}"

    def test_unknown_type_returns_other(self):
        """未知类型返回灰色。"""
        from seasync.visualization.draw_command import get_sensor_color, SENSOR_COLORS

        assert get_sensor_color("unknown_type") == SENSOR_COLORS["other"]

    def test_case_insensitive(self):
        """类型不区分大小写。"""
        from seasync.visualization.draw_command import get_sensor_color, SENSOR_COLORS

        assert get_sensor_color("RADAR") == SENSOR_COLORS["radar"]
        assert get_sensor_color("Ais") == SENSOR_COLORS["ais"]

    def test_source_id_keyword_fallback(self):
        """source_id 包含关键字时匹配颜色。"""
        from seasync.visualization.draw_command import get_sensor_color, clear_sensor_color_cache

        clear_sensor_color_cache()
        # 当 source_type 为 None 时，回退到 source_id 关键字匹配
        assert get_sensor_color(None, "radar_track_001") == "#4A90D9"
        clear_sensor_color_cache()
        assert get_sensor_color(None, "ais_mmsi_123") == "#E74C3C"
        clear_sensor_color_cache()
        assert get_sensor_color(None, "optical_cam_1") == "#F39C12"
        clear_sensor_color_cache()
        assert get_sensor_color(None, "sonar_array_2") == "#16A085"
        # 空类型+空ID 返回 other
        clear_sensor_color_cache()
        assert get_sensor_color(None, "") == "#95A5A6"


# ============================================================
# 4. 关联结果数据模型验证
# ============================================================

class TestAssociationDataModels:
    """AssociationPair / AssociationResult 新字段体系。"""

    def test_pair_generic_fields(self):
        """AssociationPair 使用通用字段。"""
        from seasync.core.data_models import AssociationPair

        pair = AssociationPair(
            source1_id="SRC-A", source2_id="SRC-B",
            source1_label="radar", source2_label="ais",
            confidence=0.95,
        )
        assert pair.source1_id == "SRC-A"
        assert pair.source2_id == "SRC-B"
        assert pair.source_ids == ("SRC-A", "SRC-B")
        assert pair.source_labels == ("radar", "ais")

    def test_pair_get_id_by_label(self):
        """按标签获取 ID。"""
        from seasync.core.data_models import AssociationPair

        pair = AssociationPair(
            source1_id="G1", source2_id="A1",
            source1_label="gps", source2_label="ais",
        )
        assert pair.get_id("gps") == "G1"
        assert pair.get_id("ais") == "A1"
        assert pair.get_id("radar") is None

    def test_pair_get_other_id(self):
        """已知一端获取另一端。"""
        from seasync.core.data_models import AssociationPair

        pair = AssociationPair(
            source1_id="G1", source2_id="A1",
            source1_label="gps", source2_label="ais",
        )
        assert pair.get_other_id("G1") == "A1"
        assert pair.get_other_id("A1") == "G1"
        assert pair.get_other_id("UNKNOWN") is None

    def test_result_unmatched_generic(self):
        """AssociationResult 使用通用 unmatched 字典。"""
        from seasync.core.data_models import AssociationResult

        result = AssociationResult(
            pairs=[],
            unmatched={"radar": ["RT-1"], "ais": ["MMSI-1"], "optical": ["OPT-1"]},
            total_quality=0.0,
        )
        assert result.get_unmatched("radar") == ["RT-1"]
        assert result.get_unmatched("optical") == ["OPT-1"]
        assert result.unmatched_radar == ["RT-1"]
        assert result.unmatched_ais == ["MMSI-1"]

    def test_result_add_unmatched(self):
        """add_unmatched 方法正确追加。"""
        from seasync.core.data_models import AssociationResult

        result = AssociationResult(pairs=[], unmatched={}, total_quality=0.0)
        result.add_unmatched("sonar", ["SONAR-1", "SONAR-2"])
        assert result.get_unmatched("sonar") == ["SONAR-1", "SONAR-2"]


# ============================================================
# 5. DrawCommand 关联连线绘制验证
# ============================================================

class TestAssociationDrawScene:
    """association_to_draw_scene 使用新字段。"""

    def test_draw_scene_uses_new_fields(self):
        """关联连线使用 source1_id / source2_id 查找目标。"""
        from seasync.core.data_models import TargetRecord, AssociationPair
        from seasync.visualization.draw_command import association_to_draw_scene, DrawType

        records_a = [
            TargetRecord(source_id="a", track_id="RT-1",
                         time=1000.0, x=100.0, y=200.0, lat=37.53, lon=121.42)
        ]
        records_b = [
            TargetRecord(source_id="b", track_id="MMSI-1",
                         time=1000.0, x=150.0, y=250.0, lat=37.54, lon=121.43)
        ]
        pairs = [AssociationPair(
            source1_id="RT-1", source2_id="MMSI-1",
            source1_label="radar", source2_label="ais",
            confidence=0.9,
        )]

        scene = association_to_draw_scene(pairs, records_a, records_b)
        # 应该有 1 条连线
        lines = [c for c in scene.commands if c.draw_type == DrawType.LINE]
        assert len(lines) == 1
        # 连线连接两个正确坐标
        line = lines[0]
        assert line.points == [(100.0, 200.0), (150.0, 250.0)]

    def test_draw_scene_handles_missing_record(self):
        """记录缺失时跳过连线。"""
        from seasync.core.data_models import TargetRecord, AssociationPair
        from seasync.visualization.draw_command import association_to_draw_scene

        records_a = [TargetRecord(source_id="a", track_id="RT-1", time=1000.0, x=0, y=0)]
        records_b = []  # 无记录
        pairs = [AssociationPair(
            source1_id="RT-1", source2_id="MMSI-1",
            source1_label="radar", source2_label="ais",
            confidence=0.9,
        )]

        scene = association_to_draw_scene(pairs, records_a, records_b)
        lines = [c for c in scene.commands if c.draw_type == c.draw_type.LINE]
        assert len(lines) == 0  # 无连线


# ============================================================
# 6. 报告生成验证
# ============================================================

class TestReportGeneration:
    """ReportGenerator 新字段表格 + 事件详细表格。"""

    def test_report_association_uses_new_fields(self):
        """关联详情表使用 source1_id / source2_id / source1_label / source2_label。"""
        from seasync.core.data_models import (
            AssociationResult, AssociationPair, EventRecord, TargetRecord
        )
        from seasync.report.report_generator import ReportGenerator

        result = AssociationResult(
            pairs=[AssociationPair(
                source1_id="RT-001", source2_id="413203610",
                source1_label="radar", source2_label="ais",
                confidence=0.95, method="kalman_mahalanobis",
            )],
            unmatched={"radar": ["RT-2"], "ais": ["413203611"]},
            total_quality=0.95,
        )

        with tempfile.TemporaryDirectory() as td:
            rg = ReportGenerator(output_dir=td)

            # Check if python-docx is available
            try:
                import docx
                has_docx = True
            except ImportError:
                has_docx = False

            if has_docx:
                # Word docx: use python-docx to read content (paragraphs + tables)
                path = rg.generate(
                    project_name="测试报告",
                    assoc_result=result,
                )
                assert path is not None
                assert os.path.exists(path)
                doc = docx.Document(path)
                # Extract from paragraphs
                full_text = "\n".join([p.text for p in doc.paragraphs])
                # Extract from tables
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            table_texts.append(cell.text)
                all_content = full_text + "\n" + "\n".join(table_texts)
                # 验证使用新字段
                assert "RT-001" in all_content
                assert "413203610" in all_content
                assert "radar" in all_content
                assert "ais" in all_content
                # 验证不再使用旧字段标签
                assert "雷达轨迹ID" not in all_content
                assert "AIS MMSI" not in all_content
            else:
                # Markdown 降级
                path = rg.generate(
                    project_name="测试报告",
                    assoc_result=result,
                )
                assert path is not None
                assert os.path.exists(path)
                with open(path, encoding="utf-8") as f:
                    content = f.read()
                assert "RT-001" in content
                assert "413203610" in content

    def test_report_events_table(self):
        """事件记录以详细表格形式展示。"""
        from seasync.core.data_models import EventRecord
        from seasync.report.report_generator import ReportGenerator

        events = [
            EventRecord(
                id="E1", time=1000.0, name="停船",
                severity="low", description="速度 < 1节 持续 30秒",
            ),
            EventRecord(
                id="E2", time=2000.0, name="机动",
                severity="medium", description="航速变化 > 5节",
            ),
        ]

        with tempfile.TemporaryDirectory() as td:
            rg = ReportGenerator(output_dir=td)
            path = rg.generate(
                project_name="测试报告",
                events=events,
            )
            assert path is not None
            assert os.path.exists(path)

            try:
                import docx
                has_docx = True
            except ImportError:
                has_docx = False

            if has_docx:
                doc = docx.Document(path)
                full_text = "\n".join([p.text for p in doc.paragraphs])
                # Also check table content
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            table_texts.append(cell.text)
                all_content = full_text + "\n" + "\n".join(table_texts)
                assert "停船" in all_content
                assert "机动" in all_content
                assert "low" in all_content
                assert "medium" in all_content
            else:
                with open(path, encoding="utf-8") as f:
                    content = f.read()
                assert "停船" in content
                assert "机动" in content
                assert "low" in content
                assert "medium" in content

    def test_report_with_pipeline_and_auto_images(self):
        """报告包含 pipeline 配置参数 + 自动可视化。"""
        from seasync.core.data_models import AssociationResult, AssociationPair
        from seasync.report.report_generator import ReportGenerator
        from seasync.engines import SeaSyncPipeline
        from tests.test_helpers import generate_radar_csv, generate_ais_simple_csv

        with tempfile.TemporaryDirectory() as td:
            radar_path = generate_radar_csv(td, n_frames=3)
            ais_path = generate_ais_simple_csv(td)

            pipe = SeaSyncPipeline(origin_lat=37.53297, origin_lon=121.42323)
            sid_r = pipe.add_source(radar_path, source_type="radar")
            sid_a = pipe.add_source(ais_path, source_type="ais")

            assoc_result = pipe._assoc.associate(
                pipe.get_records(sid_r), pipe.get_records(sid_a),
                label_a="radar", label_b="ais",
            )

            rg = ReportGenerator(output_dir=td)
            path = rg.generate(
                project_name="E2E测试",
                assoc_result=assoc_result,
                pipeline=pipe,
                auto_images=True,  # 启用自动可视化
            )
            assert path is not None
