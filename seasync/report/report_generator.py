"""
SeaSync V2.2 ReportGenerator — 专业Word报告生成器。
特性：
- 封面页（项目名称/生成时间/版本号）
- 目录
- 分节编号
- 专业表格样式
- 图表嵌入+图题
- 页眉/页脚/页码
- 无python-docx时降级Markdown
- 通用型：支持任意传感器对（雷达/AIS/光学/声纳/GPS等）
"""

# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 荣火
from __future__ import annotations
from typing import List, Dict, Any, Optional
import os
import tempfile
from datetime import datetime

from ..core.data_models import AssociationResult, EventRecord, AlignmentResult

# python-docx 相关导入（模块级，供各方法共享）
try:
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    _HAVE_DOCX = True
except ImportError:
    Pt = RGBColor = Inches = Cm = Emu = None
    WD_ALIGN_PARAGRAPH = WD_ORIENT = qn = None
    _HAVE_DOCX = False
from ..engines import SeaSyncPipeline


def _safe_str(val: Any) -> str:
    if val is None:
        return "N/A"
    if isinstance(val, float):
        return f"{val:.4f}"
    return str(val)


class ReportGenerator:
    """SeaSync 专业报告生成器。"""

    VERSION = "SeaSync V2.2"

    def __init__(self, output_dir: Optional[str] = None,
                 company_name: str = "SeaSync",
                 logo_path: Optional[str] = None) -> None:
        self.output_dir = output_dir or os.path.join(os.getcwd(), "output")
        os.makedirs(self.output_dir, exist_ok=True)
        self.company_name = company_name
        self.logo_path = logo_path

    def generate(
        self,
        project_name: str,
        assoc_result: Optional[AssociationResult] = None,
        events: Optional[List[EventRecord]] = None,
        alignment_result: Optional[AlignmentResult] = None,
        summary: Optional[Dict[str, Any]] = None,
        image_paths: Optional[List[str]] = None,
        pipeline: Optional[SeaSyncPipeline] = None,
        auto_images: bool = True,  # 新增：自动从关联结果生成可视化图片
    ) -> str:
        """生成专业Word报告，返回文件路径。"""
        try:
            from docx import Document

            doc = Document()
            self._doc = doc
            self._set_default_style()

            # 自动可视化图片生成
            resolved_image_paths: List[str] = list(image_paths or [])
            if auto_images and assoc_result:
                auto_imgs = self._generate_association_images(
                    assoc_result, pipeline, project_name
                )
                resolved_image_paths.extend(auto_imgs)

            # 封面
            self._add_cover(project_name)
            doc.add_page_break()

            # 目录占位符
            self._add_toc()

            # 各章节
            self._add_summary_section(summary or {})
            if alignment_result:
                self._add_alignment_section(alignment_result)
            if assoc_result:
                self._add_association_section(assoc_result)
            if events:
                self._add_events_section(events)
            if resolved_image_paths:
                self._add_images_section(resolved_image_paths)
            if pipeline:
                self._add_config_section(pipeline)

            self._add_footer_section()

            # 保存
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(
                self.output_dir,
                f"SeaSync_{project_name}_{timestamp}.docx"
            )
            doc.save(output_path)
            return output_path

        except ImportError:
            return self._generate_markdown(
                project_name, assoc_result, events, alignment_result,
                summary, image_paths, pipeline
            )
        except Exception as e:
            raise RuntimeError(f"报告生成失败: {e}") from e

    # ═══════════════════════════════════════════════════════════
    #  自动可视化图片生成
    # ═══════════════════════════════════════════════════════════

    def _generate_association_images(
        self,
        assoc_result: AssociationResult,
        pipeline: Optional[SeaSyncPipeline],
        project_name: str,
    ) -> List[str]:
        """从关联结果自动生成可视化图片并嵌入报告。"""
        if not pipeline:
            return []

        generated = []
        temp_dir = tempfile.gettempdir()

        try:
            # 获取所有源
            all_sids = list(pipeline._im._adapters.keys())
            if len(all_sids) < 2:
                return generated

            # 尝试生成关联可视化图
            source_records = {}
            source_labels = {}
            for sid in all_sids:
                try:
                    recs = pipeline.get_records(sid)
                    adapter = pipeline._im._adapters.get(sid)
                    src_type = getattr(adapter, 'SOURCE_TYPE', sid)
                    source_records[sid] = recs
                    source_labels[sid] = src_type
                except Exception:
                    pass

            if len(source_records) >= 2:
                sids = list(source_records.keys())
                # 生成 N 源两两关联图
                for i in range(len(sids)):
                    for j in range(i + 1, len(sids)):
                        sid_a, sid_b = sids[i], sids[j]
                        records_a = source_records[sid_a]
                        records_b = source_records[sid_b]
                        label_a = source_labels.get(sid_a, sid_a)
                        label_b = source_labels.get(sid_b, sid_b)

                        # 获取这对源的关联结果（从 pipeline 缓存或重新计算）
                        pair_result = None
                        try:
                            pair_result = pipeline._assoc.associate(
                                records_a, records_b,
                                label_a=label_a, label_b=label_b
                            )
                        except Exception:
                            pass

                        img_path = os.path.join(
                            temp_dir,
                            f"seasync_assoc_{label_a}_{label_b}_{datetime.now():%Y%m%d_%H%M%S}.png"
                        )
                        try:
                            from ..visualization.visualization import render_association
                            title = f"{label_a}-{label_b} 关联可视化 ({project_name})"
                            render_association(
                                records_a=records_a,
                                records_b=records_b,
                                assoc_result=pair_result,
                                output_path=img_path,
                                title=title,
                                label_a=label_a,
                                label_b=label_b,
                            )
                            if os.path.exists(img_path):
                                generated.append(img_path)
                        except Exception as e:
                            import logging
                            logging.getLogger("seasync").warning(
                                f"关联可视化生成失败 {label_a}-{label_b}: {e}"
                            )
        except Exception as e:
            import logging
            logging.getLogger("seasync").warning(
                f"自动可视化图片生成失败: {e}"
            )

        return generated

    # ═══════════════════════════════════════════════════════════
    #  Word 样式
    # ═══════════════════════════════════════════════════════════

    def _set_default_style(self) -> None:
        style = self._doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(6)

        # 标题样式
        for level in range(1, 4):
            hs = self._doc.styles[f"Heading {level}"]
            hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
            hs.font.bold = True

    # ═══════════════════════════════════════════════════════════
    #  封面
    # ═══════════════════════════════════════════════════════════

    def _add_cover(self, project_name: str) -> None:
        # 留白区域
        for _ in range(6):
            self._doc.add_paragraph("")

        # 主标题
        p = self._doc.add_paragraph()
        p.alignment = 1  # 居中
        run = p.add_run("SeaSync")
        run.font.size = Pt(42)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

        p2 = self._doc.add_paragraph()
        p2.alignment = 1
        run2 = p2.add_run("实海试验多源目标关联分析报告")
        run2.font.size = Pt(20)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self._doc.add_paragraph("")

        # 分隔线
        p_line = self._doc.add_paragraph()
        p_line.alignment = 1
        run_line = p_line.add_run("─" * 40)
        run_line.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
        run_line.font.size = Pt(12)

        self._doc.add_paragraph("")

        # 项目信息表
        table = self._doc.add_table(rows=4, cols=2)
        table.alignment = 1
        info = [
            ("项目名称", project_name),
            ("报告版本", self.VERSION),
            ("生成日期", datetime.now().strftime("%Y年%m月%d日 %H:%M")),
            ("生成工具", f"{self.VERSION} ReportGenerator"),
        ]
        for i, (k, v) in enumerate(info):
            c0 = table.rows[i].cells[0]
            c1 = table.rows[i].cells[1]
            c0.text = k
            c1.text = v
            for cell in [c0, c1]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
            c0.paragraphs[0].runs[0].font.bold = True
            c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

        # 底部版本号
        self._doc.add_paragraph("")
        p_v = self._doc.add_paragraph()
        p_v.alignment = 1
        run_v = p_v.add_run(f"Confidential — {self.VERSION}")
        run_v.font.size = Pt(9)
        run_v.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ═══════════════════════════════════════════════════════════
    #  目录
    # ═══════════════════════════════════════════════════════════

    def _add_toc(self) -> None:
        self._doc.add_heading("目  录", level=1)
        p = self._doc.add_paragraph("（请在 Word 中右键更新域以生成目录）")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.runs[0].font.italic = True
        # TOC 域代码
        from docx.oxml import OxmlElement
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        p2 = self._doc.add_paragraph()
        p2._element.append(fldChar1)
        p2._element.append(instrText)
        p2._element.append(fldChar2)
        p2._element.append(fldChar3)
        self._doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    #  各章节
    # ═══════════════════════════════════════════════════════════

    def _add_summary_section(self, summary: Dict[str, Any]) -> None:
        self._doc.add_heading("1  试验统计摘要", level=1)
        if not summary:
            self._doc.add_paragraph("无统计数据。")
            return
        table = self._doc.add_table(rows=len(summary) + 1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "指标"
        hdr[1].text = "数值"
        for i, (k, v) in enumerate(summary.items(), 1):
            table.rows[i].cells[0].text = str(k)
            table.rows[i].cells[1].text = str(v)

    def _add_alignment_section(self, result: AlignmentResult) -> None:
        self._doc.add_heading("2  时间对齐结果", level=1)
        self._doc.add_paragraph(
            "对数据源进行互相关时间偏移估算，"
            "以消除记录设备间的时钟差异。"
        )
        table = self._doc.add_table(rows=3, cols=2)
        table.style = "Light Grid Accent 1"
        data = [
            ("时间偏移", f"{result.offset:+.2f} 秒"),
            ("质量分数", f"{result.quality_score:.4f}"),
            ("建议操作", result.suggestion),
        ]
        for i, (k, v) in enumerate(data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

    def _add_association_section(self, result: AssociationResult) -> None:
        """关联结果章节（通用型，支持任意传感器对）。"""
        # 动态确定传感器标签
        labels = set()
        for p in result.pairs:
            labels.add(p.source1_label)
            labels.add(p.source2_label)
        label_a = sorted(labels)[0] if len(labels) >= 1 else "源A"
        label_b = sorted(labels)[1] if len(labels) >= 2 else "源B"
        title = f"{label_a}-{label_b} 关联结果"

        self._doc.add_heading(f"3  {title}", level=1)
        self._doc.add_paragraph(
            "基于卡尔曼滤波轨迹预测与马氏距离度量，"
            "采用匈牙利算法进行全局最优关联分配。"
        )
        # 概述表
        total_unmatched = sum(
            len(v) for v in result.unmatched.values()
        ) if result.unmatched else 0
        table = self._doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        data = [
            ("关联对数", str(len(result.pairs))),
            ("未匹配目标总数", str(total_unmatched)),
            ("未匹配详情", str(result.unmatched)),
            ("总质量分数", f"{result.total_quality:.4f}"),
            ("平均置信度",
             f"{sum(p.confidence for p in result.pairs)/len(result.pairs):.3f}"
             if result.pairs else "N/A"),
        ]
        for i, (k, v) in enumerate(data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

        # 详情表（使用新字段 source1_id / source2_id / source1_label / source2_label）
        if result.pairs:
            self._doc.add_heading("关联详情", level=2)
            table2 = self._doc.add_table(rows=len(result.pairs) + 1, cols=6)
            table2.style = "Light Grid Accent 1"
            hdr2 = table2.rows[0].cells
            hdr2[0].text = f"{label_a} 目标ID"
            hdr2[1].text = f"{label_a} 类型"
            hdr2[2].text = f"{label_b} 目标ID"
            hdr2[3].text = f"{label_b} 类型"
            hdr2[4].text = "置信度"
            hdr2[5].text = "匹配方法"
            for i, p in enumerate(result.pairs, 1):
                table2.rows[i].cells[0].text = p.source1_id
                table2.rows[i].cells[1].text = p.source1_label
                table2.rows[i].cells[2].text = p.source2_id
                table2.rows[i].cells[3].text = p.source2_label
                table2.rows[i].cells[4].text = f"{p.confidence:.3f}"
                table2.rows[i].cells[5].text = p.method

    def _add_events_section(self, events: List[EventRecord]) -> None:
        """事件记录章节（详细表格形式）。"""
        self._doc.add_heading(f"4  事件记录（共 {len(events)} 项）", level=1)
        if not events:
            self._doc.add_paragraph("未检测到事件。")
            return

        # 事件详细表格
        table = self._doc.add_table(rows=len(events) + 1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "时间 (s)"
        hdr[2].text = "事件类型"
        hdr[3].text = "严重程度"
        hdr[4].text = "描述"

        severity_colors = {
            "info": RGBColor(0x21, 0x96, 0xF3),
            "low": RGBColor(0x4C, 0xAF, 0x50),
            "medium": RGBColor(0xFF, 0x98, 0x00),
            "high": RGBColor(0xF4, 0x43, 0x36),
            "critical": RGBColor(0x9C, 0x27, 0xB0),
        }

        for i, e in enumerate(events, 1):
            row = table.rows[i].cells
            row[0].text = str(i)
            row[1].text = f"{e.time:.1f}"
            row[2].text = e.name
            row[3].text = e.severity.upper()
            row[4].text = e.description

            # 严重程度颜色高亮
            sev_color = severity_colors.get(e.severity.lower(), RGBColor(0x33, 0x33, 0x33))
            for r in row[3].paragraphs[0].runs:
                r.font.color.rgb = sev_color
                r.font.bold = True

        # 事件摘要
        self._doc.add_paragraph("")
        sev_summary: Dict[str, int] = {}
        for e in events:
            sev_summary[e.severity] = sev_summary.get(e.severity, 0) + 1
        summary_parts = [f"{sev}: {cnt}" for sev, cnt in sorted(sev_summary.items())]
        p = self._doc.add_paragraph(f"严重程度分布: {', '.join(summary_parts)}")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.italic = True

    def _add_images_section(self, image_paths: List[str]) -> None:
        self._doc.add_heading("5  可视化图表", level=1)
        for i, ip in enumerate(image_paths, 1):
            if not os.path.exists(ip):
                self._doc.add_paragraph(f"⚠ 图片不存在: {os.path.basename(ip)}")
                continue
            try:
                self._doc.add_picture(ip, width=Inches(5.5))
                p = self._doc.add_paragraph(f"图{i}  {os.path.basename(ip)}")
                p.alignment = 1
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.italic = True
            except Exception as e:
                self._doc.add_paragraph(f"⚠ 图片嵌入失败 {os.path.basename(ip)}: {e}")

    def _add_config_section(self, pipeline: Any) -> None:
        self._doc.add_heading("6  配置参数", level=1)
        config = pipeline.config
        table = self._doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        params = [
            ("关联距离阈值", f"{config.distance_threshold} 米"),
            ("时间窗口", f"{config.time_window_base} ± {config.time_window_tolerance} 秒"),
            ("最小置信度", f"{config.min_confidence}"),
            ("马氏距离", "启用" if config.use_mahalanobis else "禁用"),
            ("传感器不确定度", str(config.sensor_uncertainty)),
        ]
        for i, (k, v) in enumerate(params):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

    def _add_footer_section(self) -> None:
        self._doc.add_paragraph("─" * 50)
        p = self._doc.add_paragraph(
            f"本报告由 {self.VERSION} ReportGenerator 自动生成\n"
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        p.alignment = 1
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ═══════════════════════════════════════════════════════════
    #  Markdown 降级
    # ═══════════════════════════════════════════════════════════

    def _generate_markdown(
        self,
        project_name: str,
        assoc_result: Optional[AssociationResult],
        events: Optional[List[EventRecord]],
        alignment_result: Optional[AlignmentResult],
        summary: Optional[Dict[str, Any]],
        image_paths: Optional[List[str]],
        pipeline: Any = None,
    ) -> str:
        lines = [f"# SeaSync 实海试验分析报告\n\n"
                 f"**项目**: {project_name}\n"
                 f"**版本**: {self.VERSION}\n"
                 f"**生成时间**: {datetime.now():%Y-%m-%d %H:%M:%S}\n"]
        if summary:
            lines.append("\n## 1 试验统计摘要\n")
            for k, v in summary.items():
                lines.append(f"- **{k}**: {v}")
        if alignment_result:
            lines.append(
                f"\n## 2 时间对齐结果\n"
                f"- 偏移量: {alignment_result.offset:+.1f}s\n"
                f"- 质量分数: {alignment_result.quality_score:.3f}\n"
                f"- 建议: {alignment_result.suggestion}"
            )
        if assoc_result:
            # 动态标签
            labels = set()
            for p in assoc_result.pairs:
                labels.add(p.source1_label)
                labels.add(p.source2_label)
            label_a = sorted(labels)[0] if len(labels) >= 1 else "源A"
            label_b = sorted(labels)[1] if len(labels) >= 2 else "源B"
            lines.append(
                f"\n## 3 {label_a}-{label_b} 关联结果\n"
                f"- 关联对数: {len(assoc_result.pairs)}\n"
                f"- 未匹配: {dict(assoc_result.unmatched)}\n"
                f"- 总质量: {assoc_result.total_quality:.3f}\n"
            )
            if assoc_result.pairs:
                lines.append("### 关联详情\n")
                lines.append("| {0} ID | {0} 类型 | {1} ID | {1} 类型 | 置信度 | 方法 |".format(label_a, label_b))
                lines.append("|---|---|---|---|---|---|")
                for p in assoc_result.pairs:
                    lines.append(
                        f"| {p.source1_id} | {p.source1_label} | {p.source2_id} | {p.source2_label} "
                        f"| {p.confidence:.3f} | {p.method} |"
                    )
        if events:
            lines.append(f"\n## 4 事件记录（共 {len(events)} 项）\n")
            lines.append("| 时间 | 类型 | 严重程度 | 描述 |")
            lines.append("|---|---|---|---|")
            for e in events:
                lines.append(f"| {e.time:.1f}s | {e.name} | {e.severity} | {e.description} |")
        if image_paths:
            lines.append("\n## 5 可视化\n")
            for ip in image_paths:
                lines.append(f"![{os.path.basename(ip)}]({ip})")
        lines.append(f"\n---\n*由 {self.VERSION} 自动生成*")

        path = os.path.join(
            self.output_dir,
            f"SeaSync_Report_{project_name}_{datetime.now():%Y%m%d_%H%M%S}.md"
        )
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path
